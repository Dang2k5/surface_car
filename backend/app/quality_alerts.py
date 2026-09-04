from __future__ import annotations

import hashlib
import statistics
from collections import defaultdict
from datetime import UTC, datetime, timedelta
from io import BytesIO
from typing import Any, Literal

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel

from agent.services.reasoning import DeterministicReasoningService


class InspectionFinding(BaseModel):
    inspection_id: str
    thread_id: str
    vehicle_id: str
    inspected_at: str
    defect_type: str
    classified_defect_code: str
    defect_family: str
    zone_name: str
    camera_id: str
    confidence: float
    severity: str
    decision: str
    final_status: str
    recommendation_code: str
    recommendation: str
    image_url: str


class DefectAggregate(BaseModel):
    defect_type: str
    occurrence_count: int
    affected_vehicle_count: int
    zones: list[str]
    camera_ids: list[str]
    average_confidence: float
    maximum_confidence: float
    first_seen: str
    last_seen: str


class QualityAlert(BaseModel):
    id: str
    severity: str
    status: str = "OPEN"
    defect_type: str
    zone_name: str
    camera_id: str
    occurrence_count: int
    affected_vehicle_count: int
    affected_vehicle_ids: list[str]
    related_defect_codes: list[str]
    similar_code_warning: bool
    average_confidence: float
    maximum_confidence: float
    first_seen: str
    last_seen: str
    window_hours: int
    window_size: int
    consecutive_count: int
    trigger_type: str
    predicted_root_cause: str
    # Whether centroid clustering across the group's occurrences actually supports the specific
    # equipment mechanism named in predicted_root_cause, or the group only shares a zone (5 coarse
    # regions) with defects scattered across it -- PRD.md §6.1 requires root cause stay a
    # "hypothesis cần QC xác minh", never an automatic conclusion; this field is how a caller (UI,
    # report) can tell the two apart instead of trusting free text alone.
    root_cause_evidence: Literal["COORDINATE_CLUSTER_CONFIRMED", "ZONE_ONLY_UNCONFIRMED"]
    # The three checks behind root_cause_evidence, so a report/QC reviewer can see WHY it was
    # (or wasn't) confirmed instead of trusting the tag alone -- keys: coordinate_cluster (defects
    # sit within _TIGHT_CLUSTER_STDEV of each other), single_camera (every occurrence came from
    # the same camera -- cross-camera "same spot" claims are weaker, different camera rigs
    # typically frame different parts of the vehicle), severity_at_least_warning (a bare 2-vehicle
    # WATCH-tier coincidence is too weak a sample to name specific hardware over).
    root_cause_evidence_detail: dict[str, Any]
    upstream_target_shop: str
    actionable_routing_command: str
    message_en: str
    message_vi: str
    recommendation_en: str
    recommendation_vi: str
    upstream_checks_en: list[str]
    upstream_checks_vi: list[str]
    occurrences: list[InspectionFinding]
    policy_decision: dict[str, Any]
    ai_analysis: dict[str, Any]


class QualityAlertSummary(BaseModel):
    generated_at: str
    window_hours: int
    window_size: int
    watch_consecutive_threshold: int
    watch_window_threshold: int
    minimum_occurrences: int
    in_window_threshold: int
    critical_consecutive_threshold: int
    critical_window_threshold: int
    analyzed_inspections: int
    defect_breakdown: list[DefectAggregate]
    findings: list[InspectionFinding]
    alerts: list[QualityAlert]


class RepetitionAlertService:
    """Deterministic quality trend monitor over persisted LangGraph results."""

    def __init__(self, repository: Any, policy_catalog: Any, reasoning: Any) -> None:
        self.repository = repository
        self.policy_catalog = policy_catalog
        self.reasoning = reasoning
        # Trend cards must stay low-latency and can't afford one Groq call per trend per poll
        # (dashboard refetches every 10s — see SupervisorShell) — a real deterministic instance,
        # not a hopeful getattr(self.reasoning, "fallback", ...) that always missed since no
        # reasoning backend ever set that attribute, silently routing every trend card through
        # Groq and exhausting the daily token quota.
        self._trend_reasoning = DeterministicReasoningService()

    def analyze(
        self,
        *,
        window_hours: int = 24,
        window_size: int = 10,
        # Four-tier severity ladder over the SAME two deterministic signals (a same-position
        # repeat streak, or how many distinct vehicles in the window share it) — WATCH is the
        # earliest point a repeat is even a "pattern" (a single occurrence is noise, not a
        # trend); WARNING/its thresholds are unchanged from the original 2-tier logic so
        # existing callers/behavior at that boundary don't shift; CRITICAL escalates once the
        # streak or window share is more than half the monitored window, i.e. likely a
        # systemic upstream issue rather than a handful of unlucky vehicles.
        watch_consecutive_threshold: int = 2,
        watch_window_threshold: int = 2,
        minimum_occurrences: int = 3,
        in_window_threshold: int = 4,
        critical_consecutive_threshold: int = 5,
        critical_window_threshold: int = 7,
    ) -> QualityAlertSummary:
        # Keep the four thresholds monotonically ordered (WATCH <= WARNING <= CRITICAL) even if
        # a caller passes inconsistent query params — otherwise a lower CRITICAL bound than
        # WARNING would make the WARNING tier unreachable for that signal.
        watch_consecutive_threshold = min(watch_consecutive_threshold, minimum_occurrences)
        watch_window_threshold = min(watch_window_threshold, in_window_threshold)
        critical_consecutive_threshold = max(critical_consecutive_threshold, minimum_occurrences)
        critical_window_threshold = max(critical_window_threshold, in_window_threshold)
        now = datetime.now(UTC)
        cutoff = now - timedelta(hours=window_hours)
        candidates = []
        # list_with_metadata() is ordered updated_at DESC (newest first), so once a row falls
        # before the window cutoff every remaining row is also older -- stop scanning instead
        # of touching the rest of a potentially large audit history on every call (this runs
        # every 2s off the SSE loop in backend/app/v1_api.py). `limit=500` is a defense-in-depth
        # cap, generous relative to window_size/window_hours' realistic values.
        for state in self.repository.list_with_metadata(limit=500):
            persisted_at = _parse_timestamp(state.get("_persisted_at"))
            if persisted_at < cutoff:
                break
            if not state.get("defect_detected"):
                continue
            if state.get("defect_type") not in {"scratch", "dent"}:
                continue
            if state.get("decision") == "PASS":
                continue
            candidates.append(state)

        candidates.sort(
            key=lambda item: _parse_timestamp(item.get("_persisted_at")),
            reverse=True,
        )
        records = candidates[:window_size]

        grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
        for state in records:
            key = (
                str(state.get("defect_type", "unknown")),
                _zone_name(state),
            )
            grouped[key].append(state)

        findings = sorted(
            (_finding_from_state(state) for state in records),
            key=lambda item: item.inspected_at,
            reverse=True,
        )
        defect_breakdown = _build_defect_breakdown(records)

        alerts: list[QualityAlert] = []
        for (defect_type, zone_name), items in grouped.items():
            vehicle_ids = sorted({str(item.get("vehicle_id", "UNKNOWN")) for item in items})
            consecutive_count = _leading_consecutive_count(records, defect_type, zone_name)
            window_count = len(vehicle_ids)
            severity, trigger_type = _severity_for(
                consecutive_count=consecutive_count,
                window_count=window_count,
                watch_consecutive_threshold=watch_consecutive_threshold,
                watch_window_threshold=watch_window_threshold,
                warning_consecutive_threshold=minimum_occurrences,
                warning_window_threshold=in_window_threshold,
                critical_consecutive_threshold=critical_consecutive_threshold,
                critical_window_threshold=critical_window_threshold,
            )
            if severity is None:
                continue
            camera_ids = sorted({str(item.get("camera_id") or "unknown_camera") for item in items})
            camera_id = camera_ids[0] if len(camera_ids) == 1 else "MULTI_CAMERA"
            confidences = [float(item.get("confidence") or 0) for item in items]
            related_codes = sorted(
                {str(item.get("classified_defect_code")) for item in items if item.get("classified_defect_code")}
            )
            similar_code_warning = len(related_codes) > 1
            timestamps = sorted(_parse_timestamp(item.get("_persisted_at")) for item in items)
            root_cause, target_shop, root_cause_evidence, root_cause_evidence_detail = _predicted_root_cause(
                defect_type, items, severity=severity, camera_id=camera_id
            )
            routing_command = "ROUTE_AFFECTED_BATCH_TO_OFFLINE_INSPECTION_BUFFER"
            key_text = f"{defect_type}|{zone_name}|{window_size}"
            alert_id = hashlib.sha256(key_text.encode("utf-8")).hexdigest()[:16]
            trend_state = {
                "defect_type": defect_type,
                "confidence": round(sum(confidences) / len(confidences), 4),
                "zone_name": zone_name,
                "camera_id": camera_id,
                "severity": severity,
                "evidence_tags": [
                    "affected_vehicle_list",
                    "time_window",
                    "camera_and_zone_group",
                ],
                "trend_context": {
                    "affected_vehicle_count": len(vehicle_ids),
                    "affected_vehicle_ids": vehicle_ids,
                    "window_hours": window_hours,
                    "window_size": window_size,
                    "consecutive_count": consecutive_count,
                    "trigger_type": trigger_type,
                },
            }
            policy = self.policy_catalog.evaluate_named("FNS-TREND-001", trend_state)
            # Dashboard trend aggregation must remain low-latency. LLM reasoning
            # is reserved for an explicit inspection; trend cards use the same
            # deterministic policy fallback instead of issuing N sequential API calls.
            analysis = self._trend_reasoning.analyze(trend_state, policy)
            alerts.append(
                QualityAlert(
                    id=alert_id,
                    severity=severity,
                    defect_type=defect_type,
                    zone_name=zone_name,
                    camera_id=camera_id,
                    occurrence_count=len(items),
                    affected_vehicle_count=len(vehicle_ids),
                    affected_vehicle_ids=vehicle_ids,
                    related_defect_codes=related_codes,
                    similar_code_warning=similar_code_warning,
                    average_confidence=round(sum(confidences) / len(confidences), 4),
                    maximum_confidence=round(max(confidences), 4),
                    first_seen=timestamps[0].isoformat(),
                    last_seen=timestamps[-1].isoformat(),
                    window_hours=window_hours,
                    window_size=window_size,
                    consecutive_count=consecutive_count,
                    trigger_type=trigger_type,
                    predicted_root_cause=root_cause,
                    root_cause_evidence=root_cause_evidence,
                    root_cause_evidence_detail=root_cause_evidence_detail,
                    upstream_target_shop=target_shop,
                    actionable_routing_command=routing_command,
                    message_en=(
                        f"Repeated {defect_type} detections were found in {zone_name}: "
                        f"{consecutive_count} consecutive and {len(vehicle_ids)}/{window_size} recent vehicles. "
                        f"Related QC codes: {', '.join(related_codes) or 'unclassified'}."
                    ),
                    message_vi=(
                        f"Phát hiện lỗi {defect_type} lặp lại tại {zone_name}: "
                        f"{consecutive_count} xe liên tiếp và {len(vehicle_ids)}/{window_size} xe gần nhất. "
                        f"Mã lỗi liên quan: {', '.join(related_codes) or 'chưa phân loại'}."
                    ),
                    recommendation_en=(
                        "Keep affected vehicles controlled, confirm the trend with a named QC reviewer, "
                        "and escalate to the upstream process owner before release."
                    ),
                    recommendation_vi=(
                        "Kiểm soát các xe liên quan, yêu cầu QC xác nhận xu hướng và thông báo chủ công đoạn "
                        "phía trước trước khi cho phép release."
                    ),
                    upstream_checks_en=_upstream_checks("en", zone_name, camera_id),
                    upstream_checks_vi=_upstream_checks("vi", zone_name, camera_id),
                    occurrences=sorted(
                        (_finding_from_state(item) for item in items),
                        key=lambda item: item.inspected_at,
                        reverse=True,
                    ),
                    policy_decision=policy.model_dump(mode="json"),
                    ai_analysis=analysis.model_dump(mode="json"),
                )
            )
        severity_rank = {"CRITICAL": 0, "WARNING": 1, "WATCH": 2}
        alerts.sort(
            key=lambda item: (severity_rank.get(item.severity, 3), -item.affected_vehicle_count)
        )
        return QualityAlertSummary(
            generated_at=now.isoformat(),
            window_hours=window_hours,
            window_size=window_size,
            watch_consecutive_threshold=watch_consecutive_threshold,
            watch_window_threshold=watch_window_threshold,
            minimum_occurrences=minimum_occurrences,
            in_window_threshold=in_window_threshold,
            critical_consecutive_threshold=critical_consecutive_threshold,
            critical_window_threshold=critical_window_threshold,
            analyzed_inspections=len(records),
            defect_breakdown=defect_breakdown,
            findings=findings,
            alerts=alerts,
        )


def _finding_from_state(state: dict[str, Any]) -> InspectionFinding:
    return InspectionFinding(
        inspection_id=str(state.get("inspection_id") or "UNKNOWN"),
        thread_id=str(state.get("thread_id") or "UNKNOWN"),
        vehicle_id=str(state.get("vehicle_id") or "UNKNOWN"),
        inspected_at=_parse_timestamp(state.get("_persisted_at")).isoformat(),
        defect_type=str(state.get("defect_type") or "unknown"),
        classified_defect_code=str(state.get("classified_defect_code") or ""),
        defect_family=str(state.get("defect_family") or ""),
        zone_name=_zone_name(state),
        camera_id=str(state.get("camera_id") or "unknown_camera"),
        confidence=round(float(state.get("confidence") or 0), 4),
        severity=str(state.get("severity") or "UNASSESSED"),
        decision=str(state.get("decision") or "UNKNOWN"),
        final_status=str(state.get("final_status") or "UNKNOWN"),
        recommendation_code=str(state.get("recommendation_code") or ""),
        recommendation=str(state.get("recommendation") or ""),
        image_url=str(state.get("image_url") or ""),
    )


def _zone_name(state: dict[str, Any]) -> str:
    return str(state.get("zone_name") or "unknown_zone")


def _leading_consecutive_count(
    records: list[dict[str, Any]],
    defect_type: str,
    zone_name: str,
) -> int:
    count = 0
    for state in records:
        if str(state.get("defect_type")) != defect_type or _zone_name(state) != zone_name:
            break
        count += 1
    return count


def _severity_for(
    *,
    consecutive_count: int,
    window_count: int,
    watch_consecutive_threshold: int,
    watch_window_threshold: int,
    warning_consecutive_threshold: int,
    warning_window_threshold: int,
    critical_consecutive_threshold: int,
    critical_window_threshold: int,
) -> tuple[str | None, str | None]:
    """Rank the same (defect_type, zone_name) group's two deterministic repeat signals —
    a same-position streak across consecutive vehicles, and how many distinct vehicles in
    the monitored window share it — against an escalating ladder shared by both signals.
    CONSECUTIVE is reported whenever it alone reaches the tier, even if WINDOW_FREQUENCY also
    qualifies: a same-position streak is the stronger signal of a live upstream cause (the
    same defect keeps recurring right now), whereas a same window count from vehicles at
    scattered ranks doesn't come with that same immediacy."""
    if (
        consecutive_count >= critical_consecutive_threshold
        or window_count >= critical_window_threshold
    ):
        trigger = (
            "CONSECUTIVE"
            if consecutive_count >= critical_consecutive_threshold
            else "WINDOW_FREQUENCY"
        )
        return "CRITICAL", trigger
    if (
        consecutive_count >= warning_consecutive_threshold
        or window_count >= warning_window_threshold
    ):
        trigger = (
            "CONSECUTIVE"
            if consecutive_count >= warning_consecutive_threshold
            else "WINDOW_FREQUENCY"
        )
        return "WARNING", trigger
    if consecutive_count >= watch_consecutive_threshold or window_count >= watch_window_threshold:
        trigger = (
            "CONSECUTIVE" if consecutive_count >= watch_consecutive_threshold else "WINDOW_FREQUENCY"
        )
        return "WATCH", trigger
    return None, None


# Same defect+zone repeating at a near-identical frame position across vehicles is the actual
# evidence PRD.md §6.1 ties a specific equipment hypothesis to ("cùng tọa độ" for dents, "cùng
# đường kẻ dọc" for scratches) -- zone_name alone is only 5 coarse body regions, so two vehicles
# can share a zone with defects nowhere near the same spot. A stdev this tight (~8% of frame
# width/height) across every occurrence's primary detection center is what tells the two apart.
_TIGHT_CLUSTER_STDEV = 0.08


def _primary_center(state: dict[str, Any]) -> tuple[float, float] | None:
    detections = state.get("detections") or []
    if not detections:
        return None
    primary_id = state.get("primary_detection_id")
    primary = next(
        (item for item in detections if item.get("detection_id") == primary_id),
        detections[0],
    )
    measurements = primary.get("visual_measurements") or {}
    center_x = measurements.get("center_x_ratio")
    center_y = measurements.get("center_y_ratio")
    if center_x is None or center_y is None:
        return None
    return float(center_x), float(center_y)


def _is_tight_cluster(centers: list[tuple[float, float]]) -> bool:
    if len(centers) < 2:
        return False
    return (
        statistics.pstdev(c[0] for c in centers) <= _TIGHT_CLUSTER_STDEV
        and statistics.pstdev(c[1] for c in centers) <= _TIGHT_CLUSTER_STDEV
    )


def _predicted_root_cause(
    defect_type: str,
    items: list[dict[str, Any]],
    *,
    severity: str,
    camera_id: str,
) -> tuple[
    str,
    str,
    Literal["COORDINATE_CLUSTER_CONFIRMED", "ZONE_ONLY_UNCONFIRMED"],
    dict[str, Any],
]:
    """Root cause is a hypothesis for QC to verify, never an automatic conclusion (PRD.md §6.1).
    Only name a specific equipment mechanism when ALL three independent signals support it --
    each rules out a different way the group could look like a real repeat but not be one:
    - coordinate_cluster: occurrences actually sit at the same frame position, not just the same
      coarse zone (a zone is 5 body regions; two vehicles can share one with defects nowhere near
      the same spot).
    - single_camera: every occurrence came from the same camera. A "same spot" claim spanning
      different camera rigs is weaker -- different cameras typically frame different parts of the
      vehicle, so cross-camera agreement is more likely coincidence than a shared physical cause.
    - severity_at_least_warning: a bare WATCH-tier alert can be as few as 2 vehicles -- too small
      a sample to send a maintenance team after specific hardware over.
    Any one of the three failing means the group doesn't earn a confident, named mechanism --
    it gets the generic "not enough evidence yet" hypothesis instead."""
    centers = [center for center in (_primary_center(item) for item in items) if center is not None]
    detail = {
        "coordinate_cluster": _is_tight_cluster(centers),
        "single_camera": camera_id != "MULTI_CAMERA",
        "severity_at_least_warning": severity != "WATCH",
        "occurrence_count": len(items),
    }
    confirmed = detail["coordinate_cluster"] and detail["single_camera"] and detail["severity_at_least_warning"]
    evidence: Literal["COORDINATE_CLUSTER_CONFIRMED", "ZONE_ONLY_UNCONFIRMED"] = (
        "COORDINATE_CLUSTER_CONFIRMED" if confirmed else "ZONE_ONLY_UNCONFIRMED"
    )
    if defect_type == "dent":
        shop = "Stamping / Body Shop"
        if confirmed:
            root_cause = (
                "Vết móp lặp lại tại cùng một tọa độ, cùng camera, trên đủ số xe liên tiếp để "
                "loại trừ trùng hợp ngẫu nhiên — giả thuyết: khuôn dập (stamping die) dính "
                "bavia/mạt kim loại hoặc tay gắp robot bị kẹt dị vật đúng vị trí đó. Cần QC xác "
                "minh trực tiếp thiết bị trước khi kết luận."
            )
        else:
            root_cause = (
                "Nhiều xe cùng bị móp trong cùng vùng kiểm tra nhưng chưa đủ bằng chứng (vị trí "
                "lỗi không tập trung rõ rệt, dữ liệu đến từ nhiều camera khác nhau, hoặc số lần "
                "lặp lại còn ít) để quy về một cơ chế thiết bị cụ thể; QC cần xác minh thêm "
                "(khuôn dập, tay gắp, hoặc va chạm rời rạc) trước khi kết luận."
            )
        return root_cause, shop, evidence, detail
    shop = "Body / Paint Handling Process"
    if confirmed:
        root_cause = (
            "Vết xước lặp lại cùng một đường kẻ/tọa độ, cùng camera, trên đủ số xe liên tiếp để "
            "loại trừ trùng hợp ngẫu nhiên — giả thuyết: con lăn băng tải hoặc thanh dẫn hướng bị "
            "cọ xát/mòn đúng vị trí đó. Cần QC xác minh trực tiếp thiết bị trước khi kết luận."
        )
    else:
        root_cause = (
            "Nhiều xe cùng bị xước trong cùng vùng kiểm tra nhưng chưa đủ bằng chứng (vị trí lỗi "
            "không tập trung rõ rệt, dữ liệu đến từ nhiều camera khác nhau, hoặc số lần lặp lại "
            "còn ít) để quy về một cơ chế thiết bị cụ thể; QC cần xác minh thêm trước khi kết luận."
        )
    return root_cause, shop, evidence, detail


def _build_defect_breakdown(records: list[dict[str, Any]]) -> list[DefectAggregate]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for state in records:
        grouped[str(state.get("defect_type") or "unknown")].append(state)

    result: list[DefectAggregate] = []
    for defect_type, items in grouped.items():
        confidences = [float(item.get("confidence") or 0) for item in items]
        timestamps = sorted(_parse_timestamp(item.get("_persisted_at")) for item in items)
        result.append(
            DefectAggregate(
                defect_type=defect_type,
                occurrence_count=len(items),
                affected_vehicle_count=len({str(item.get("vehicle_id") or "UNKNOWN") for item in items}),
                zones=sorted({_zone_name(item) for item in items}),
                camera_ids=sorted({str(item.get("camera_id") or "unknown_camera") for item in items}),
                average_confidence=round(sum(confidences) / len(confidences), 4),
                maximum_confidence=round(max(confidences), 4),
                first_seen=timestamps[0].isoformat(),
                last_seen=timestamps[-1].isoformat(),
            )
        )
    return sorted(result, key=lambda item: (-item.occurrence_count, item.defect_type))


def _parse_timestamp(value: Any) -> datetime:
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=UTC)
    except (TypeError, ValueError):
        return datetime.now(UTC)


def _upstream_checks(language: str, zone_name: str, camera_id: str) -> list[str]:
    if language == "vi":
        return [
            f"Xác nhận camera {camera_id}: tiêu cự, ánh sáng, vị trí gá và độ sạch của lens.",
            f"Kiểm tra đồ gá, dụng cụ tiếp xúc và thao tác có thể ảnh hưởng tới {zone_name}.",
            "Đối chiếu thời điểm lỗi với ca sản xuất và lịch sử bảo trì thiết bị.",
            "Cách ly mẫu liên quan và thực hiện kiểm tra xác nhận dưới ánh sáng kiểm soát.",
            "Ghi nhận người phụ trách, hành động khắc phục và tiêu chí release sau kiểm tra lại.",
        ]
    return [
        f"Verify {camera_id}: focus, lighting, fixture position, and lens cleanliness.",
        f"Inspect fixtures, contact tools, and handling operations that may affect {zone_name}.",
        "Correlate occurrence times with shift and equipment maintenance history.",
        "Contain affected samples and perform confirmation inspection under controlled lighting.",
        "Record the owner, corrective action, and release criteria after reinspection.",
    ]


def build_quality_alert_report(summary: QualityAlertSummary) -> BytesIO:
    """Create an operator-ready DOCX using the compact_reference_guide preset."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "VISUAL QC AGENT  |  FNS QUALITY TREND ALERT"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string("607B8D")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Internal QC working report  |  Generated by deterministic trend rules").font.size = Pt(8)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("QUALITY ESCALATION REPORT")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("C24D5A")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Repeated Defect Alert")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("0B2940")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Agent-generated signal for upstream process verification")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("607B8D")

    metadata = document.add_table(rows=2, cols=4)
    set_table_geometry(metadata, [1350, 3330, 1350, 3330])
    metadata_values = [
        ("Generated", _format_dt(summary.generated_at), "Window", f"{summary.window_hours} hours"),
        ("Inspections", str(summary.analyzed_inspections), "Open alerts", str(len(summary.alerts))),
    ]
    for row, values in zip(metadata.rows, metadata_values):
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if index % 2 == 0:
                shade_cell(cell, "EAF1F5")
                cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    document.add_heading("Agent assessment", level=1)
    lead = document.add_paragraph()
    lead.paragraph_format.space_after = Pt(8)
    if summary.alerts:
        critical = sum(item.severity == "CRITICAL" for item in summary.alerts)
        lead.add_run(
            f"The trend monitor found {len(summary.alerts)} repeated-defect group(s), including "
            f"{critical} critical alert(s). QC must verify the upstream process before release."
        ).bold = True
    else:
        lead.add_run("No repeated-defect group crossed the configured threshold in this window.").bold = True

    document.add_heading("Defect summary from inspection history", level=1)
    if summary.defect_breakdown:
        defect_table = document.add_table(rows=1, cols=5)
        set_table_geometry(defect_table, [1500, 1000, 1000, 4000, 1860])
        for cell, value in zip(
            defect_table.rows[0].cells,
            ("Defect", "Events", "Vehicles", "Zones / cameras", "Confidence"),
        ):
            cell.text = value
            shade_cell(cell, "DCEAF1")
            cell.paragraphs[0].runs[0].bold = True
        for item in summary.defect_breakdown:
            row = defect_table.add_row()
            values = (
                item.defect_type,
                str(item.occurrence_count),
                str(item.affected_vehicle_count),
                f"{', '.join(item.zones)}\n{', '.join(item.camera_ids)}",
                f"Avg {item.average_confidence:.1%}\nMax {item.maximum_confidence:.1%}",
            )
            for cell, value in zip(row.cells, values):
                cell.text = value
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
    else:
        document.add_paragraph("No defect finding is retained in the selected monitoring window.")

    document.add_heading("Inspection-level findings", level=1)
    if summary.findings:
        finding_table = document.add_table(rows=1, cols=7)
        set_table_geometry(finding_table, [1250, 1100, 1100, 1700, 850, 1050, 2310])
        for cell, value in zip(
            finding_table.rows[0].cells,
            ("Time", "Inspection", "Vehicle", "Defect / location", "Conf.", "Severity", "Decision / route"),
        ):
            cell.text = value
            shade_cell(cell, "DCEAF1")
            cell.paragraphs[0].runs[0].bold = True
        for finding in summary.findings:
            row = finding_table.add_row()
            values = (
                _format_dt(finding.inspected_at),
                finding.inspection_id,
                finding.vehicle_id,
                f"{finding.defect_type}\n{finding.zone_name} / {finding.camera_id}",
                f"{finding.confidence:.1%}",
                finding.severity,
                f"{finding.decision}\n{finding.final_status}\n{finding.recommendation_code}",
            )
            for cell, value in zip(row.cells, values):
                cell.text = value
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(7.5)
    else:
        document.add_paragraph("No inspection-level finding is available for this report.")

    for index, alert in enumerate(summary.alerts, start=1):
        document.add_heading(f"Alert {index}: {alert.defect_type.upper()} - {alert.zone_name}", level=2)
        alert_table = document.add_table(rows=5, cols=2)
        set_table_geometry(alert_table, [2160, 7200])
        rows = [
            ("Severity", alert.severity),
            ("Detection source", alert.camera_id),
            ("Affected vehicles", f"{alert.affected_vehicle_count} ({', '.join(alert.affected_vehicle_ids)})"),
            ("Confidence", f"Average {alert.average_confidence:.1%} | Maximum {alert.maximum_confidence:.1%}"),
            ("Observed window", f"{_format_dt(alert.first_seen)} to {_format_dt(alert.last_seen)}"),
        ]
        for row, (label, value) in zip(alert_table.rows, rows):
            row.cells[0].text = label
            row.cells[1].text = value
            shade_cell(row.cells[0], "EAF1F5")
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        document.add_heading("Inspections contributing to this alert", level=3)
        occurrence_table = document.add_table(rows=1, cols=6)
        set_table_geometry(occurrence_table, [1450, 1450, 1450, 1550, 1100, 2360])
        for cell, value in zip(
            occurrence_table.rows[0].cells,
            ("Time", "Inspection", "Vehicle", "Confidence", "Severity", "Final route"),
        ):
            cell.text = value
            shade_cell(cell, "F3E4E7")
            cell.paragraphs[0].runs[0].bold = True
        for occurrence in alert.occurrences:
            row = occurrence_table.add_row()
            values = (
                _format_dt(occurrence.inspected_at),
                occurrence.inspection_id,
                occurrence.vehicle_id,
                f"{occurrence.confidence:.1%}",
                occurrence.severity,
                occurrence.final_status,
            )
            for cell, value in zip(row.cells, values):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)

        document.add_heading("Required upstream checks", level=3)
        for check in alert.upstream_checks_en:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.2
            paragraph.add_run(check)
        recommendation = document.add_paragraph()
        recommendation.paragraph_format.space_before = Pt(6)
        recommendation.paragraph_format.space_after = Pt(10)
        label = recommendation.add_run("Disposition: ")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string("A73343")
        recommendation.add_run(alert.recommendation_en)

        document.add_heading("Policy and reasoning provenance", level=3)
        provenance = document.add_paragraph()
        policy = alert.policy_decision
        analysis = alert.ai_analysis
        provenance.add_run(
            f"Policy: {policy.get('policy_id')} @ {policy.get('policy_revision')} "
            f"({policy.get('policy_status')})\n"
        ).bold = True
        provenance.add_run(
            f"Reasoning: {analysis.get('provider')} / {analysis.get('model')}\n"
            f"Analysis: {analysis.get('summary_en')}"
        )
        for reference in policy.get("references", []):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(
                f"{reference.get('id')}: {reference.get('title')} — {reference.get('url')}"
            )

    document.add_heading("QC sign-off", level=1)
    signoff = document.add_table(rows=3, cols=2)
    set_table_geometry(signoff, [2700, 6660])
    for row, (label, value) in zip(
        signoff.rows,
        (("QC reviewer", ""), ("Upstream process owner", ""), ("Corrective action / release evidence", "")),
    ):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], "EAF1F5")
        row.cells[0].paragraphs[0].runs[0].bold = True

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _format_dt(value: str) -> str:
    return _parse_timestamp(value).astimezone(UTC).strftime("%Y-%m-%d %H:%M UTC")


def shade_cell(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table: Any, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_properties = cell._tc.get_or_add_tcPr()
            tc_width = tc_properties.get_or_add_tcW()
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width))
            margins = tc_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_properties.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
