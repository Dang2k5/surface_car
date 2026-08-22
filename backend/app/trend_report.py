from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .quality_alerts import set_table_geometry, shade_cell

GROUP_LABELS = {"hour": "Giờ", "day": "Ngày", "shift": "Ca", "lot": "Lô"}


def build_trend_report(
    rows: list[dict[str, Any]],
    *,
    group_by: str,
    shift_id: str | None,
    lot_id: str | None,
    station_id: str | None,
    date_from: str | None,
    date_to: str | None,
) -> BytesIO:
    """Operator-ready DOCX for GET /api/trend, scoped by day/shift/lot/station.

    Separate from build_quality_alert_report (agent/services/quality_alerts.py),
    which only covers repeated-defect alerts — this covers routine production
    quality reporting for a chosen date range / lot / station / shift.
    """
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Heading 1", 15, "2E74B5"),):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "VISUAL QC AGENT  |  BÁO CÁO XU HƯỚNG CHẤT LƯỢNG"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string("607B8D")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(
        f"Generated {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')} | Visual QC Agent"
    ).font.size = Pt(8)

    title = document.add_paragraph()
    title_run = title.add_run("Báo cáo chất lượng sản xuất")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string("0B2940")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(f"Nhóm theo: {GROUP_LABELS.get(group_by, group_by)}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("607B8D")

    filters = document.add_table(rows=1, cols=2)
    set_table_geometry(filters, [2500, 6860])
    filters.rows[0].cells[0].text = "Bộ lọc"
    filters.rows[0].cells[1].text = "Giá trị"
    shade_cell(filters.rows[0].cells[0], "DCEAF1")
    shade_cell(filters.rows[0].cells[1], "DCEAF1")
    for label, value in (
        ("Khoảng ngày", f"{date_from or '—'} → {date_to or '—'}"),
        ("Ca làm việc", shift_id or "Tất cả"),
        ("Lô sản xuất", lot_id or "Tất cả"),
        ("Trạm", station_id or "Tất cả"),
    ):
        row = filters.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        shade_cell(row.cells[0], "EAF1F5")

    total_inspections = sum(r["total_inspections"] for r in rows)
    total_pass = sum(r["pass_count"] for r in rows)
    total_fail = sum(r["fail_count"] for r in rows)
    overall_pass_rate = round((total_pass / total_inspections) * 100, 1) if total_inspections else 0.0

    document.add_heading("Tổng quan", level=1)
    summary = document.add_table(rows=2, cols=4)
    set_table_geometry(summary, [2340, 2340, 2340, 2340])
    for cell, value in zip(summary.rows[0].cells, ("Tổng inspection", "PASS", "FAIL", "Tỷ lệ đạt")):
        cell.text = value
        shade_cell(cell, "DCEAF1")
    for cell, value in zip(
        summary.rows[1].cells,
        (str(total_inspections), str(total_pass), str(total_fail), f"{overall_pass_rate}%"),
    ):
        cell.text = value

    document.add_heading(f"Chi tiết theo {GROUP_LABELS.get(group_by, group_by).lower()}", level=1)
    if rows:
        table = document.add_table(rows=1, cols=7)
        set_table_geometry(table, [1500, 1300, 1300, 1300, 1300, 1300, 1360])
        for cell, value in zip(
            table.rows[0].cells,
            (
                GROUP_LABELS.get(group_by, group_by),
                "Tổng",
                "Trầy xước",
                "Móp",
                "PASS",
                "FAIL",
                "Tỷ lệ lỗi",
            ),
        ):
            cell.text = value
            shade_cell(cell, "DCEAF1")
        for row_data in rows:
            row = table.add_row()
            values = (
                str(row_data["group_value"]),
                str(row_data["total_inspections"]),
                str(row_data["scratch_count"]),
                str(row_data["dent_count"]),
                str(row_data["pass_count"]),
                str(row_data["fail_count"]),
                f"{row_data['pass_fail_rate']}%",
            )
            for cell, value in zip(row.cells, values):
                cell.text = value
    else:
        document.add_paragraph("Không có dữ liệu inspection nào khớp với bộ lọc đã chọn.")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
