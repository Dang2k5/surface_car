from __future__ import annotations

from io import BytesIO

from pydantic import BaseModel, Field


class DocumentTextExtractionError(ValueError):
    """Raised when the uploaded file cannot be read as text (corrupt, unsupported, or empty)."""


def extract_text(data: bytes, filename: str) -> str:
    """Best-effort text extraction for the document types QC supervisors upload
    (PDF work instructions, DOCX control plans). Raises DocumentTextExtractionError
    when nothing usable comes out, so the caller can fail the request with a clear
    reason instead of asking an LLM to extract a policy from an empty string."""
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix == "pdf":
        text = _extract_pdf_text(data)
    elif suffix == "docx":
        text = _extract_docx_text(data)
    else:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError as error:
            raise DocumentTextExtractionError(
                f"Unsupported file type: .{suffix or 'unknown'}. Upload a PDF, DOCX, or UTF-8 text file."
            ) from error
    text = text.strip()
    if len(text) < 40:
        raise DocumentTextExtractionError(
            "Extracted text is too short to contain a usable QC policy "
            "(scanned/image-only PDFs are not supported — this pipeline is text-only)."
        )
    return text


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as error:
        raise DocumentTextExtractionError(f"Could not parse PDF: {error}") from error
    return "\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    try:
        document = Document(BytesIO(data))
    except Exception as error:
        raise DocumentTextExtractionError(f"Could not parse DOCX: {error}") from error
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


class PolicyDraft(BaseModel):
    """Shape mirrors backend.app.qc_schemas.PolicyItemCreate closely enough for the
    frontend to load it straight into the existing create/edit form (rules.tsx) —
    the QC supervisor still reviews and edits every field before anything is saved."""

    suggested_id: str = Field(description="Proposed policy id, e.g. FNS-SURFACE-002")
    title: str
    defect_types: list[str] = Field(default_factory=list)
    conditions: list[str] = Field(default_factory=list)
    required_evidence: list[str] = Field(default_factory=list)
    steps: list[str] = Field(default_factory=list)
    action_code: str
    final_status: str
    test_drive_allowed: bool | None = None
    human_required: bool = False


class PolicySourceDraft(BaseModel):
    """Metadata for the new `sources` catalog entry that will point back at the
    uploaded document once the policy is saved."""

    document_family: str
    revision: str
    title: str
    section: str
    effective_date: str | None = None


class PolicyExtractionResult(BaseModel):
    policy: PolicyDraft
    source: PolicySourceDraft
    extraction_notes_vi: str = Field(
        default="", description="Caveats the QC supervisor should double-check before approving."
    )
    provider: str = "groq"
    model: str = ""
